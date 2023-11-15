import docx
from docx.shared import Inches, Pt, Length
import sys
import docx2txt
import re

#dictionary for key words to look for and replace in template file
replace_dict = {
    "addressed_to": sys.argv[2]
}

# file path to template given as command line arguments
template_file_path = sys.argv[1]

# open connection to Word Document
template_doc = docx.Document(template_file_path)

# read in each paragraph in file
doc_contents = [p.text for p in template_doc.paragraphs]

new_doc_contents = []
for paragraph in doc_contents:
    paragraph_content = paragraph.split(" ")
    # print(paragraph_content)
    new_paragraph = ""
    new_word = ""
    for word in paragraph_content:
        keyword = re.search(r"\{([A-Za-z0-9_]+)\}", word)
        if keyword and keyword.group(1) in replace_dict:
            #NOTE: I'm personally adding an apostrphe s to the end of the word
            new_paragraph += (replace_dict[keyword.group(1)])
        else:
            new_paragraph += word
        new_paragraph += " "
    new_doc_contents.append(new_paragraph)


new_document = docx.Document()
for p in range(len(new_doc_contents)):
    paragraph = new_doc_contents[p]
    para = new_document.add_paragraph(paragraph)
    run = para.add_run()
    style = new_document.styles['Normal']
    font = style.font
    #change font size and name here if you desire
    font.name = 'Amasis MT Pro'
    font.size = docx.shared.Pt(12)
    if p > 0 and p < len(new_doc_contents) - 2:
        para.paragraph_format.first_line_indent = Inches(0.5)
    para.paragraph_format.line_spacing = 1

new_document.save(sys.argv[2] + ".docx")

